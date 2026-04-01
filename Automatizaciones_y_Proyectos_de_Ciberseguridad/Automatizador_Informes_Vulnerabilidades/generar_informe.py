#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
from datetime import datetime, timedelta
from pathlib import Path
import sys
import os
import json
import time
import requests
import re
import html
import traceback
import functools
import logging
import yaml

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START

# Intentar importar deep-translator (reemplaza googletrans)
try:
    from deep_translator import GoogleTranslator
    TRANSLATOR_DISPONIBLE = True
except ImportError:
    TRANSLATOR_DISPONIBLE = False

# Configurar matplotlib para no usar interfaz gráfica
matplotlib.use('Agg')

# DIRECTORIO BASE (siempre junto al script)

SCRIPT_DIR = Path(os.path.dirname(os.path.abspath(__file__)))

# LOGGING
log_path = SCRIPT_DIR / "informe_generator.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(log_path, encoding="utf-8"),
    ],
)
logger = logging.getLogger(__name__)


# CONFIG YAML
def cargar_config():
    config_path = SCRIPT_DIR / "config.yaml"
    defaults = {
        "cache": {"expiry_days": 30, "file": "cves_cache.json"},
        "report": {"top_n_hosts": 5, "top_n_cves": 5, "output_dir": "."},
        "charts": {
            "colors": {
                "critical": "#d32f2f",
                "high": "#ff9800",
                "medium": "#fdd835",
                "low": "#4caf50",
            }
        },
        "translation": {"enabled": True, "target_language": "es"},
    }
    if config_path.exists():
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                user_cfg = yaml.safe_load(f) or {}
            # Merge profundo sencillo: nivel 1 y nivel 2
            for section, values in user_cfg.items():
                if section in defaults and isinstance(defaults[section], dict):
                    defaults[section].update(values)
                else:
                    defaults[section] = values
            logger.info("config.yaml cargado correctamente.")
        except Exception as e:
            logger.warning(f"No se pudo leer config.yaml, usando valores por defecto: {e}")
    else:
        logger.warning("config.yaml no encontrado. Usando valores por defecto.")
    return defaults

CONFIG = cargar_config()

# API KEY NVD (config.yaml tiene prioridad; variable de entorno como fallback)
NVD_API_KEY = CONFIG.get("nvd_api_key") or os.getenv("NVD_API_KEY") or None
if NVD_API_KEY:
    logger.info("API Key de NVD detectada.")
else:
    logger.warning("Sin API Key de NVD - usando limite de peticiones publicas.")


# FUNCIONES AUXILIARES
def escapar_caracteres_xml(texto):
    """Escapa caracteres que pueden romper el XML de Word."""
    if not texto:
        return ""
    return (
        texto.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace('"', "&quot;")
             .replace("'", "&apos;")
    )


def traducir_a_español(texto):
    """Traduce texto al español usando deep-translator."""
    if not TRANSLATOR_DISPONIBLE or not texto or len(texto) < 5:
        return texto
    if not CONFIG["translation"]["enabled"]:
        return texto
    try:
        target = CONFIG["translation"]["target_language"]
        traduccion = GoogleTranslator(source="auto", target=target).translate(texto)
        return traduccion if traduccion else texto
    except Exception as e:
        logger.warning(f"Traducción fallida: {e}")
        return texto


def log_tiempo_ejecucion(func):
    """Decorador para medir tiempos y capturar errores."""
    @functools.wraps(func)
    def wrapper(self, *args, **kwargs):
        inicio = time.time()
        try:
            resultado = func(self, *args, **kwargs)
            duracion = time.time() - inicio
            if hasattr(self, "metricas_red"):
                self.metricas_red["tiempo_total_consultas"] += duracion
            if duracion > 2:
                logger.info(f"{func.__name__} tardó {duracion:.2f}s")
            return resultado
        except Exception as e:
            duracion = time.time() - inicio
            if hasattr(self, "metricas_red"):
                self.metricas_red["errores_red"] += 1
                self.metricas_red["tiempo_total_consultas"] += duracion
            logger.error(f"Error en {func.__name__}: {e}")
            return None
    return wrapper


# CLASE: BuscadorCVERobusto

class BuscadorCVERobusto:
    """Busca detalles de CVEs con caché y reintentos."""

    def __init__(self):
        self.cache_file = SCRIPT_DIR / CONFIG["cache"]["file"]
        self.cache_expiracion_dias = CONFIG["cache"]["expiry_days"]
        self.cache = self._cargar_cache()
        self.tiene_api_key = NVD_API_KEY is not None

        self.metricas_red = {
            "tiempo_total_consultas": 0.0,
            "exitos_nvd": 0,
            "fallbacks_cvedetails": 0,
            "fallbacks_generico": 0,
            "errores_red": 0,
        }
        self._limpiar_cache_expirado()

    def _cargar_cache(self):
        if self.cache_file.exists():
            try:
                with open(self.cache_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"No se pudo cargar caché: {e}")
        return {}

    def _limpiar_cache_expirado(self):
        if not self.cache:
            return
        ahora = datetime.now()
        limite = ahora - timedelta(days=self.cache_expiracion_dias)
        cache_limpio = {}

        for cve_id, info in self.cache.items():
            try:
                if "timestamp" in info and datetime.fromisoformat(info["timestamp"]) >= limite:
                    cache_limpio[cve_id] = info
            except Exception as e:
                logger.debug(f"Entrada de caché inválida para {cve_id}: {e}")

        if len(cache_limpio) < len(self.cache):
            removed = len(self.cache) - len(cache_limpio)
            logger.info(f"Caché: {removed} entradas expiradas eliminadas.")
            self.cache = cache_limpio
            self._guardar_cache()

    def _guardar_cache(self):
        try:
            with open(self.cache_file, "w", encoding="utf-8") as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"No se pudo guardar caché: {e}")

    def _sanitizar_html(self, texto):
        if not texto:
            return ""
        texto = html.unescape(texto)
        texto = re.sub(r"<br\s*/?>", " ", texto, flags=re.IGNORECASE)
        texto = re.sub(r"<[^>]+>", "", texto)
        return re.sub(r"\s+", " ", texto).strip()

    def _limpiar_source(self, source_raw):
        """Limpia el identificador de fuente (ej: secure@microsoft.com -> Microsoft)."""
        if not source_raw:
            return "N/A"
        if "@" in source_raw:
            try:
                dominio = source_raw.split("@")[1]
                nombre = dominio.split(".")[0]
                if nombre == "mitre":
                    return "MITRE"
                if nombre == "nist":
                    return "NIST"
                return nombre.capitalize()
            except Exception:
                return source_raw
        return source_raw

    def _extraer_cvss(self, vuln):
        """Extrae score y vector CVSS (v3.1 > v3.0 > v2.0)."""
        score = None
        vector = None
        version_usada = None

        metrics = vuln.get("metrics", {})
        for version_key, label in [
            ("cvssMetricV31", "3.1"),
            ("cvssMetricV30", "3.0"),
            ("cvssMetricV2", "2.0"),
        ]:
            if version_key in metrics and metrics[version_key]:
                try:
                    cvss_data = metrics[version_key][0].get("cvssData", {})
                    score = cvss_data.get("baseScore")
                    vector = cvss_data.get("vectorString")
                    version_usada = label
                    break
                except Exception as e:
                    logger.debug(f"Error extrayendo CVSS {label}: {e}")

        return score, vector, version_usada

    @log_tiempo_ejecucion
    def _buscar_nvd(self, cve_id):
        """Busca en NVD y devuelve dict con descripcion, source y cvss."""
        url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}"
        headers = {"User-Agent": "Mozilla/5.0"}
        if self.tiene_api_key:
            headers["apiKey"] = NVD_API_KEY

        for intento in range(3):
            try:
                response = requests.get(url, headers=headers, timeout=15)
                if response.status_code == 200:
                    data = response.json()
                    if data.get("vulnerabilities"):
                        vuln = data["vulnerabilities"][0]["cve"]

                        source_raw = vuln.get("sourceIdentifier", "NVD")
                        source_clean = self._limpiar_source(source_raw)

                        desc_final = ""
                        for desc in vuln.get("descriptions", []):
                            if desc["lang"] == "es":
                                desc_final = self._sanitizar_html(desc["value"])
                                break
                        if not desc_final and vuln.get("descriptions"):
                            desc_final = self._sanitizar_html(vuln["descriptions"][0]["value"])

                        cvss_score, cvss_vector, cvss_version = self._extraer_cvss(vuln)

                        if desc_final:
                            self.metricas_red["exitos_nvd"] += 1
                            return {
                                "detalle": desc_final,
                                "source": source_clean,
                                "cvss_score": cvss_score,
                                "cvss_vector": cvss_vector,
                                "cvss_version": cvss_version,
                            }
                return None

            except requests.exceptions.Timeout:
                espera = 0.5 * (2 ** intento)
                logger.warning(f"Timeout buscando {cve_id}, reintento {intento+1} en {espera}s")
                time.sleep(espera)
            except requests.exceptions.RequestException as e:
                logger.warning(f"Error de red buscando {cve_id}: {e}")
                time.sleep(0.5)

        return None

    @log_tiempo_ejecucion
    def _buscar_cve_details(self, cve_id):
        """Fallback a CVE Details."""
        # CORRECCIÓN v2: URL correcta (era https.www. en v1)
        url = f"https://www.cvedetails.com/cve/{cve_id}/"
        for intento in range(3):
            try:
                response = requests.get(
                    url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10
                )
                if response.status_code == 200 and len(response.text) > 1000:
                    match = re.search(
                        r"<b>Description:</b>\s*</td>\s*<td[^>]*>\s*([^<]+)",
                        response.text,
                    )
                    if match:
                        self.metricas_red["fallbacks_cvedetails"] += 1
                        return self._sanitizar_html(match.group(1).strip())
            except requests.exceptions.RequestException as e:
                logger.debug(f"CVE Details fallback falló para {cve_id}: {e}")
                time.sleep(0.5)
        return None

    def _generar_descripcion_generica(self, cve_id, package_name=""):
        patrones = {
            "firefox": "Vulnerabilidad de seguridad en Mozilla Firefox.",
            "teams": "Vulnerabilidad en Microsoft Teams.",
            "chrome": "Vulnerabilidad de seguridad en Google Chrome.",
            "edge": "Vulnerabilidad en Microsoft Edge.",
            "office": "Vulnerabilidad en Microsoft Office.",
            "windows": "Vulnerabilidad de seguridad del sistema Windows.",
            "java": "Vulnerabilidad en Java.",
        }
        package_lower = package_name.lower() if package_name else ""
        for patron, desc in patrones.items():
            if patron in package_lower:
                self.metricas_red["fallbacks_generico"] += 1
                return desc
        self.metricas_red["fallbacks_generico"] += 1
        return "Vulnerabilidad de seguridad identificada."

    def buscar_cve_nvd(self, cve_id, package_name=""):
        """Busca CVE: Caché -> NVD -> CVE Details -> Genérico."""

        # 1. Caché (la traducción ya está almacenada en detalle_es)
        if cve_id in self.cache:
            return self.cache[cve_id]

        info = {
            "cve_id": cve_id,
            "detalle": "",
            "detalle_es": "",        # v2: traducción almacenada en caché
            "origen": "",
            "source": "N/A",
            "cvss_score": None,      # v2: puntuación CVSS
            "cvss_vector": None,     # v2: vector CVSS
            "cvss_version": None,    # v2: versión CVSS usada
            "url": f"https://nvd.nist.gov/vuln/detail/{cve_id}",
            "timestamp": datetime.now().isoformat(),
        }

        # 2. NVD
        resultado_nvd = self._buscar_nvd(cve_id)
        if resultado_nvd:
            info["detalle"] = resultado_nvd["detalle"]
            info["detalle_es"] = traducir_a_español(resultado_nvd["detalle"])
            info["source"] = resultado_nvd["source"]
            info["cvss_score"] = resultado_nvd["cvss_score"]
            info["cvss_vector"] = resultado_nvd["cvss_vector"]
            info["cvss_version"] = resultado_nvd["cvss_version"]
            info["origen"] = "NVD"
        else:
            # 3. Fallback CVE Details
            detalle = self._buscar_cve_details(cve_id)
            if detalle:
                info["detalle"] = detalle
                info["detalle_es"] = traducir_a_español(detalle)
                info["origen"] = "CVE Details"
                info["source"] = "CVE Details"
            else:
                # 4. Genérico
                detalle = self._generar_descripcion_generica(cve_id, package_name)
                info["detalle"] = detalle
                info["detalle_es"] = detalle   # ya en español
                info["origen"] = "Genérico"
                info["source"] = "Automático"

        self.cache[cve_id] = info
        self._guardar_cache()
        return info


# CLASE: GeneradorInformeDocx

class GeneradorInformeDocx:
    def __init__(self, ruta_excel, infraestructura="Sistema", mes_anio=None):
        self.ruta_excel = ruta_excel
        self.infraestructura = infraestructura
        self.mes_anio = mes_anio if mes_anio else datetime.now().strftime("%B %Y")

        self.df = None
        self.metricas = {}
        self.rutas_graficos = []
        self.buscador = BuscadorCVERobusto()

        self.top_n_hosts = CONFIG["report"]["top_n_hosts"]
        self.top_n_cves  = CONFIG["report"]["top_n_cves"]
        self.colores      = CONFIG["charts"]["colors"]

    # Carga y validación
    COLUMNAS_ESPERADAS = ["timestamp", "host", "cve", "severity", "package_name", "package_version", "status"]

    def cargar_datos(self):
        logger.info(f"Cargando datos: {self.ruta_excel}")
        try:
            df = pd.read_excel(self.ruta_excel, skiprows=1)
            if len(df.columns) < len(self.COLUMNAS_ESPERADAS):
                raise ValueError(
                    f"El Excel tiene {len(df.columns)} columnas, se esperan al menos {len(self.COLUMNAS_ESPERADAS)}."
                )
            df.columns = self.COLUMNAS_ESPERADAS + list(df.columns[len(self.COLUMNAS_ESPERADAS):])
            self.df = df
            logger.info(f"{len(self.df)} registros cargados.")
        except FileNotFoundError:
            logger.error(f"Archivo no encontrado: {self.ruta_excel}")
            sys.exit(1)
        except ValueError as e:
            logger.error(f"Estructura del Excel inválida: {e}")
            sys.exit(1)
        except Exception as e:
            logger.error(f"Fallo al leer Excel: {e}")
            sys.exit(1)

    # Métricas 

    def calcular_metricas(self):
        logger.info("Calculando métricas...")
        self.metricas["total_vulnerabilities"] = len(self.df)
        self.metricas["unique_cves"]  = self.df["cve"].nunique()
        self.metricas["unique_hosts"] = self.df["host"].nunique()

        severity_counts = self.df["severity"].value_counts()
        for sev in ["critical", "high", "medium", "low"]:
            key = sev.capitalize()
            self.metricas[sev] = int(severity_counts.get(key, 0))

        total = self.metricas["total_vulnerabilities"]
        for sev in ["critical", "high", "medium", "low"]:
            count = self.metricas[sev]
            self.metricas[f"{sev}_pct"] = round((count / total * 100), 2) if total > 0 else 0

    def generar_top_hosts(self):
        logger.info("Generando top hosts...")
        top_hosts = (
            self.df.groupby("host").size().sort_values(ascending=False).head(self.top_n_hosts)
        )
        datos = []
        for host in top_hosts.index:
            h_df = self.df[self.df["host"] == host]
            datos.append({
                "Host":    host,
                "Total":   len(h_df),
                "Críticas": len(h_df[h_df["severity"] == "Critical"]),
                "Altas":   len(h_df[h_df["severity"] == "High"]),
                "Medias":  len(h_df[h_df["severity"] == "Medium"]),
                "Bajas":   len(h_df[h_df["severity"] == "Low"]),
            })
        self.metricas["top_hosts_data"] = pd.DataFrame(datos)

    def generar_top_cves(self):
        logger.info("Generando top CVEs...")
        top_cves = (
            self.df.groupby("cve").size().sort_values(ascending=False).head(self.top_n_cves)
        )
        datos = []
        for cve in top_cves.index:
            c_df = self.df[self.df["cve"] == cve]
            datos.append({
                "CVE":        cve,
                "Ocurrencias": len(c_df),
                "Severidad":  c_df["severity"].mode()[0],
                "Hosts":      c_df["host"].nunique(),
            })
        self.metricas["top_cves_data"] = pd.DataFrame(datos)

    # Gráficos 

    def generar_graficos(self):
        logger.info("Generando gráficos...")
        temp_dir = SCRIPT_DIR / "temp_graficos"
        temp_dir.mkdir(exist_ok=True)
        self.rutas_graficos = []

        col = self.colores

        # 1. Pastel Severidad
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = ["Crítica", "Alta", "Media", "Baja"]
        sizes  = [self.metricas["critical"], self.metricas["high"],
                  self.metricas["medium"],  self.metricas["low"]]
        colors = [col["critical"], col["high"], col["medium"], col["low"]]

        datos = [(l, s, c) for l, s, c in zip(labels, sizes, colors) if s > 0]
        if datos:
            l, s, c = zip(*datos)
            total = sum(s)
            new_labels = [f"{lbl} ({sz/total*100:.1f}%)" for lbl, sz in zip(l, s)]
            ax.pie(s, labels=new_labels, colors=c, startangle=90,
                   wedgeprops={"edgecolor": "black", "linewidth": 1.5},
                   textprops={"fontweight": "bold", "fontsize": 11})
            ax.set_title("Distribución por Severidad", fontsize=14, fontweight="bold", pad=30)

        plt.tight_layout()
        ruta = temp_dir / "distribucion.png"
        plt.savefig(ruta, dpi=300, bbox_inches="tight")
        plt.close()
        self.rutas_graficos.append(str(ruta))

        # 2. Top Hosts
        fig, ax = plt.subplots(figsize=(10, 6))
        top_h = self.df.groupby("host").size().sort_values(ascending=False).head(self.top_n_hosts)
        bars = ax.barh(range(len(top_h)), top_h.values, color="#1976d2")
        ax.set_yticks(range(len(top_h)))
        ax.set_yticklabels(top_h.index)
        ax.set_xlabel("Vulnerabilidades", fontsize=12, fontweight="bold")
        ax.set_title(f"Top {self.top_n_hosts} Hosts", fontsize=14, fontweight="bold")
        for bar in bars:
            ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2,
                    f" {int(bar.get_width())}", va="center", fontweight="bold")
        plt.tight_layout()
        ruta = temp_dir / "hosts.png"
        plt.savefig(ruta, dpi=300, bbox_inches="tight")
        plt.close()
        self.rutas_graficos.append(str(ruta))

        # 3. Top CVEs
        fig, ax = plt.subplots(figsize=(10, 6))
        top_c = self.df.groupby("cve").size().sort_values(ascending=False).head(self.top_n_cves)
        bars = ax.barh(range(len(top_c)), top_c.values, color="#f57c00")
        ax.set_yticks(range(len(top_c)))
        ax.set_yticklabels(top_c.index, fontsize=10)
        ax.set_xlabel("Ocurrencias", fontsize=12, fontweight="bold")
        ax.set_title(f"Top {self.top_n_cves} CVEs Más Frecuentes", fontsize=14, fontweight="bold")
        for bar in bars:
            ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2,
                    f" {int(bar.get_width())}", va="center", fontweight="bold")
        plt.tight_layout()
        ruta = temp_dir / "cves.png"
        plt.savefig(ruta, dpi=300, bbox_inches="tight")
        plt.close()
        self.rutas_graficos.append(str(ruta))

    # Detalle CVEs 

    def _procesar_lista_cves(self, severidad, prefijo):
        df_sev = self.df[self.df["severity"] == severidad]
        cves_unicos = df_sev["cve"].unique()
        lista = []

        if len(cves_unicos) > 0:
            logger.info(f"Procesando {len(cves_unicos)} CVEs {severidad}...")
            pausa = 6 if not self.buscador.tiene_api_key else 0.3

            for idx, cve in enumerate(cves_unicos):
                logger.info(f"  [{prefijo}-{idx+1}/{len(cves_unicos)}] {cve}...")
                cve_df = df_sev[df_sev["cve"] == cve]
                pkg = cve_df["package_name"].iloc[0] if len(cve_df) > 0 else ""

                info = self.buscador.buscar_cve_nvd(cve, pkg)

                lista.append({
                    "cve":           cve,
                    "occurrences":   len(cve_df),
                    "hosts_affected": cve_df["host"].nunique(),
                    "package_name":  pkg,
                    "info":          info,
                })
                time.sleep(pausa)
        return lista

    def generar_detalles_vulnerabilidades(self):
        logger.info("Buscando detalles de CVEs...")
        self.metricas["top_critical_cves"] = self._procesar_lista_cves("Critical", "C")
        self.metricas["top_high_cves"]     = self._procesar_lista_cves("High",     "H")
        self.metricas["top_medium_cves"]   = self._procesar_lista_cves("Medium",   "M")
        self.metricas["top_low_cves"]      = self._procesar_lista_cves("Low",      "L")

    #  Generación DOCX 

    def _add_table_with_data(self, doc, headers, data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for row_d in data:
            row_cells = table.add_row().cells
            for i, v in enumerate(row_d):
                row_cells[i].text = str(v)

    def _agregar_cve_con_hipervinculo(self, doc, item):
        info = item["info"]
        detalle_mostrar = info.get("detalle_es") or info.get("detalle", "")

        # Título CVE
        p = doc.add_paragraph()
        run = p.add_run(info["cve_id"])
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)

        # CVSS Score (v2: nuevo campo)
        if info.get("cvss_score") is not None:
            cvss_text = f"CVSS {info['cvss_version']}: {info['cvss_score']}"
            if info.get("cvss_vector"):
                cvss_text += f"  ({info['cvss_vector']})"
            doc.add_paragraph(cvss_text)

        p_det = doc.add_paragraph()
        p_det.add_run("Detalle técnico: ").bold = True
        p_det.add_run(detalle_mostrar)
        p_oc = doc.add_paragraph()
        p_oc.add_run("Ocurrencias: ").bold = True
        p_oc.add_run(str(item["occurrences"]))
        p_ha = doc.add_paragraph()
        p_ha.add_run("Hosts afectados: ").bold = True
        p_ha.add_run(str(item["hosts_affected"]))
        p_pk = doc.add_paragraph()
        p_pk.add_run("Paquete: ").bold = True
        p_pk.add_run(str(item["package_name"]))

        if info.get("source") and info["source"] != "N/A":
            p_src = doc.add_paragraph()
            p_src.add_run("Fuente: ").bold = True
            p_src.add_run(str(info["source"]))

        # Enlace (Field Codes)
        p_ref = doc.add_paragraph()
        p_ref.add_run("Referencia: ").bold = True

        try:
            url = info["url"].replace('"', "%22")

            run_begin = p_ref.add_run()
            fldChar_begin = OxmlElement("w:fldChar")
            fldChar_begin.set(qn("w:fldCharType"), "begin")
            run_begin._r.append(fldChar_begin)

            run_instr = p_ref.add_run()
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = f'HYPERLINK "{url}"'
            run_instr._r.append(instrText)

            run_sep = p_ref.add_run()
            fldChar_sep = OxmlElement("w:fldChar")
            fldChar_sep.set(qn("w:fldCharType"), "separate")
            run_sep._r.append(fldChar_sep)

            run_disp = p_ref.add_run(info["url"])
            run_disp.font.color.rgb = RGBColor(0, 0, 255)
            run_disp.font.underline = True

            run_end = p_ref.add_run()
            fldChar_end = OxmlElement("w:fldChar")
            fldChar_end.set(qn("w:fldCharType"), "end")
            run_end._r.append(fldChar_end)

        except Exception as e:
            logger.warning(f"No se pudo crear hipervínculo para {info['cve_id']}: {e}")
            p_ref.add_run(info["url"])

        doc.add_paragraph()

    def _configurar_toc(self, doc):
        try:
            settings = doc.settings.element
            update = OxmlElement("w:updateFields")
            update.set(qn("w:val"), "true")
            settings.append(update)
        except Exception as e:
            logger.debug(f"No se pudo configurar TOC: {e}")

    def generar_docx(self, nombre_salida=None):
        logger.info("Generando DOCX...")

        if nombre_salida is None:
            periodo_safe  = self.mes_anio.replace("/", "-").replace("\\", "-")
            empresa_safe  = self.infraestructura.replace("/", "-").replace("\\", "-")
            nombre_salida = f"INFORME_VULNERABILIDADES_{empresa_safe}_{periodo_safe}.docx"

        output_dir = Path(CONFIG["report"]["output_dir"])
        output_dir.mkdir(parents=True, exist_ok=True)
        ruta_final = output_dir / nombre_salida

        doc = Document()

        # 1. PORTADA
        section = doc.sections[0]
        section.top_margin    = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin   = Inches(0)
        section.right_margin  = Inches(0)
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)

        portada_path = SCRIPT_DIR / "imagenes" / "portada.png"
        if portada_path.exists():
            try:
                doc.add_picture(str(portada_path), width=Inches(8.27))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"No se pudo insertar portada: {e}")
                doc.add_heading("INFORME", 0)
        else:
            doc.add_heading("INFORME DE VULNERABILIDADES", 0)

        doc.add_section(WD_SECTION_START.NEW_PAGE)

        # 2. ÍNDICE
        section = doc.sections[-1]
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.footer.is_linked_to_previous = False

        t = doc.add_paragraph("INFORME DE VULNERABILIDADES")
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].font.size = Pt(16)
        t.runs[0].bold      = True
        t.runs[0].underline = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Empresa: ").bold = True
        p.runs[0].font.size = Pt(16)
        p.add_run(self.infraestructura).font.size = Pt(16)

        p2 = doc.add_paragraph(f"{self.mes_anio}")
        p2.alignment       = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(14)
        p2.runs[0].font.bold = True

        doc.add_paragraph()

        pi = doc.add_paragraph()
        pi.add_run("Índice").font.size = Pt(18)
        pi.runs[0].bold = True

        p = doc.add_paragraph()
        r = p.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        r._r.append(fldChar)
        instr = OxmlElement("w:instrText")
        instr.text = 'TOC \\o "1-1" \\h \\z \\u'
        r._r.append(instr)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r._r.append(fldChar2)

        doc.add_section(WD_SECTION_START.NEW_PAGE)

        # 3. CONTENIDO
        section = doc.sections[-1]
        section.footer.is_linked_to_previous = False

        sectPr = section._sectPr
        pgNumType = OxmlElement("w:pgNumType")
        pgNumType.set(qn("w:start"), "1")
        sectPr.append(pgNumType)

        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_foot.add_run("Página ")

        r = p_foot.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        r._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        r._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        r._r.append(fld2)

        logo_path = SCRIPT_DIR / "imagenes" / "logo.png"
        if logo_path.exists():
            try:
                pl  = footer.add_paragraph()
                pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = pl.add_run()
                run.add_picture(str(logo_path), width=Inches(0.7))
            except Exception as e:
                logger.debug(f"No se pudo insertar logo: {e}")

        # Resumen ejecutivo
        doc.add_heading("Resumen Ejecutivo", 1)
        doc.add_paragraph(
            f"Durante el análisis de {self.mes_anio}, se identificaron "
            f"{self.metricas['total_vulnerabilities']} vulnerabilidades en "
            f"{self.metricas['unique_hosts']} hosts. "
            f"Críticas: {self.metricas['critical']}, Altas: {self.metricas['high']}."
        )

        doc.add_paragraph()
        doc.add_heading("Métricas Generales", 2)
        metrics = [
            ["Total",       str(self.metricas["total_vulnerabilities"])],
            ["CVEs Únicos", str(self.metricas["unique_cves"])],
            ["Hosts",       str(self.metricas["unique_hosts"])],
            ["Críticas",    f"{self.metricas['critical']} ({self.metricas['critical_pct']}%)"],
            ["Altas",       f"{self.metricas['high']} ({self.metricas['high_pct']}%)"],
            ["Medias",      f"{self.metricas['medium']} ({self.metricas['medium_pct']}%)"],
            ["Bajas",       f"{self.metricas['low']} ({self.metricas['low_pct']}%)"],
        ]
        self._add_table_with_data(doc, ["Métrica", "Valor"], metrics)
        doc.add_page_break()

        doc.add_heading("Análisis por Severidad", 1)
        doc.add_paragraph("Distribución de vulnerabilidades por nivel de riesgo.")
        doc.add_paragraph()
        if self.rutas_graficos:
            doc.add_picture(self.rutas_graficos[0], width=Inches(5.5))
        doc.add_page_break()

        doc.add_heading("Análisis de Infraestructura", 1)
        doc.add_heading(f"Top {self.top_n_hosts} Hosts", 2)

        hosts_data = []
        for _, row in self.metricas["top_hosts_data"].iterrows():
            hosts_data.append([
                row["Host"], str(row["Total"]), str(row["Críticas"]),
                str(row["Altas"]), str(row["Medias"]), str(row["Bajas"])
            ])
        self._add_table_with_data(doc, ["Host", "Tot", "Crit", "Alt", "Med", "Baj"], hosts_data)

        doc.add_paragraph()
        if len(self.rutas_graficos) > 1:
            doc.add_picture(self.rutas_graficos[1], width=Inches(5.5))
        doc.add_page_break()

        doc.add_heading("Análisis de CVEs", 1)
        doc.add_heading(f"Top {self.top_n_cves} CVEs", 2)

        cves_data = []
        for _, row in self.metricas["top_cves_data"].iterrows():
            cves_data.append([row["CVE"], str(row["Ocurrencias"]), row["Severidad"], str(row["Hosts"])])
        self._add_table_with_data(doc, ["CVE", "Ocur", "Sev", "Hosts"], cves_data)

        doc.add_paragraph()
        if len(self.rutas_graficos) > 2:
            doc.add_picture(self.rutas_graficos[2], width=Inches(5.5))

        # Detalles por severidad
        for sev, lista in [
            ("Críticas", self.metricas["top_critical_cves"]),
            ("Altas",    self.metricas["top_high_cves"]),
            ("Medias",   self.metricas["top_medium_cves"]),
            ("Bajas",    self.metricas["top_low_cves"]),
        ]:
            if lista:
                doc.add_page_break()
                doc.add_heading(f"Vulnerabilidades {sev}", 1)
                for item in lista:
                    self._agregar_cve_con_hipervinculo(doc, item)

        self._configurar_toc(doc)

        try:
            doc.save(str(ruta_final))
            logger.info(f"Generado: {ruta_final}")
        except PermissionError:
            logger.error(f"No se pudo guardar. Cierra el archivo '{nombre_salida}' e inténtalo de nuevo.")
            return None

        # Limpieza temporales
        try:
            for ruta_g in self.rutas_graficos:
                os.remove(ruta_g)
            os.rmdir(str(SCRIPT_DIR / "temp_graficos"))
        except Exception:
            pass

        return str(ruta_final)

    def ejecutar_completo(self):
        logger.info("=" * 70)
        logger.info("  GENERADOR DE INFORMES DE VULNERABILIDADES v2")
        logger.info("=" * 70)

        self.cargar_datos()
        self.calcular_metricas()
        self.generar_top_hosts()
        self.generar_top_cves()
        self.generar_detalles_vulnerabilidades()
        self.generar_graficos()
        self.generar_docx()

        logger.info("=" * 70)
        logger.info("[FIN] Proceso completado.")
        logger.info(
            f"Métricas NVD: éxitos={self.buscador.metricas_red['exitos_nvd']}, "
            f"fallback_details={self.buscador.metricas_red['fallbacks_cvedetails']}, "
            f"genérico={self.buscador.metricas_red['fallbacks_generico']}, "
            f"errores={self.buscador.metricas_red['errores_red']}"
        )
        logger.info("=" * 70)

# MAIN

def main():
    if len(sys.argv) < 2:
        print("Uso: python generar_informe.py <excel> [empresa] [periodo]")
        print("  Ejemplo: python generar_informe.py datos.xlsx \"Acme Corp\" \"Marzo 2026\"")
        sys.exit(1)

    generador = GeneradorInformeDocx(
        sys.argv[1],
        sys.argv[2] if len(sys.argv) > 2 else "Empresa",
        sys.argv[3] if len(sys.argv) > 3 else None,
    )
    generador.ejecutar_completo()


if __name__ == "__main__":
    main()
